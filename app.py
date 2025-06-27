from flask import Flask, render_template, request, send_file, redirect, url_for, flash, jsonify
import os
import tempfile
import whisper
from docx import Document
import math
from pyannote.audio import Pipeline
from dotenv import load_dotenv

app = Flask(__name__)
app.secret_key = 'supersecretkey'
UPLOAD_FOLDER = tempfile.gettempdir()
ALLOWED_EXTENSIONS = {'mp3', 'wav', 'm4a', 'ogg', 'flac'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Cargar modelo Whisper una sola vez
model = whisper.load_model('base')
# Cargar pipeline de diarización de pyannote (requiere token HuggingFace)
load_dotenv()
HUGGINGFACE_TOKEN = os.environ.get('HUGGINGFACE_TOKEN')
pipeline = Pipeline.from_pretrained('pyannote/speaker-diarization-3.1', use_auth_token=HUGGINGFACE_TOKEN) if HUGGINGFACE_TOKEN else None

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/', methods=['GET'])
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload():
    if 'audiofile' not in request.files:
        return jsonify({'success': False, 'message': 'No se seleccionó ningún archivo'}), 400
    file = request.files['audiofile']
    if file.filename == '':
        return jsonify({'success': False, 'message': 'No se seleccionó ningún archivo'}), 400
    if file and allowed_file(file.filename):
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
        file.save(filepath)
        if os.path.getsize(filepath) == 0:
            return jsonify({'success': False, 'message': 'El archivo está vacío. Por favor, selecciona un archivo de audio válido.'}), 400
        # Validación extra: intentar abrir el archivo con ffmpeg para asegurarse que es audio válido
        import subprocess
        try:
            result = subprocess.run([
                'ffmpeg', '-v', 'error', '-i', filepath, '-f', 'null', '-'
            ], capture_output=True, text=True)
            if result.stderr:
                return jsonify({'success': False, 'message': 'El archivo no es un audio válido o está dañado.'}), 400
        except Exception:
            return jsonify({'success': False, 'message': 'No se pudo validar el archivo de audio.'}), 400
        return jsonify({'success': True, 'filename': file.filename, 'message': 'Archivo subido exitosamente.'})
    else:
        return jsonify({'success': False, 'message': 'Formato de archivo no permitido'}), 400

@app.route('/transcribe', methods=['POST'])
def transcribe():
    data = request.get_json()
    filename = data.get('filename')
    idioma = data.get('idioma', 'es')
    num_speakers = int(data.get('num_speakers', 2))
    franjas = data.get('franjas', False)
    if not filename:
        return jsonify({'success': False, 'message': 'No se recibió el nombre del archivo.'}), 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return jsonify({'success': False, 'message': 'El archivo no existe en el servidor.'}), 400
    try:
        # Diarización con pyannote
        if pipeline is None:
            return jsonify({'success': False, 'message': 'pyannote.audio requiere un HuggingFace Token. Por favor, configura la variable de entorno HUGGINGFACE_TOKEN.'}), 500
        diarization = pipeline({'audio': filepath}, num_speakers=num_speakers)
        # Transcripción con Whisper
        result = model.transcribe(filepath, fp16=False, language=idioma, word_timestamps=True)
        segments = result.get('segments', [])
        # Asignar segmentos a hablantes
        speaker_segments = []
        for seg in segments:
            # Buscar el hablante que más se solapa con el segmento
            best_speaker = None
            best_overlap = 0
            for turn in diarization.itertracks(yield_label=True):
                start, end, speaker = turn[0][0], turn[0][1], turn[2]
                overlap = max(0, min(seg['end'], end) - max(seg['start'], start))
                if overlap > best_overlap:
                    best_overlap = overlap
                    best_speaker = speaker
            speaker_segments.append({'speaker': best_speaker, 'text': seg['text'], 'start': seg['start'], 'end': seg['end']})
        # Agrupar por franjas si el usuario lo desea
        text_content = ''
        if franjas:
            grouped = []
            current = []
            current_start = 0
            for seg in speaker_segments:
                if seg['start'] - current_start >= 600 and current:
                    grouped.append(current)
                    current = []
                    current_start = seg['start']
                current.append(seg)
            if current:
                grouped.append(current)
            for i, group in enumerate(grouped):
                start_min = math.floor(group[0]['start'] / 60)
                end_min = math.ceil(group[-1]['end'] / 60)
                text_content += f"\n--- Franja {i+1}: minuto {start_min} a {end_min} ---\n"
                for seg in group:
                    text_content += f"{seg['speaker']}: {seg['text']}\n"
        else:
            for seg in speaker_segments:
                text_content += f"{seg['speaker']}: {seg['text']}\n"
        # Guardar TXT
        txt_path = filepath + '.txt'
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write(text_content)
        # Guardar DOCX
        doc = Document()
        doc.add_heading('Transcripción de audio', 0)
        if franjas:
            for i, group in enumerate(grouped):
                start_min = math.floor(group[0]['start'] / 60)
                end_min = math.ceil(group[-1]['end'] / 60)
                doc.add_heading(f'Franja {i+1}: minuto {start_min} a {end_min}', level=1)
                for seg in group:
                    doc.add_paragraph(f"{seg['speaker']}: {seg['text']}")
        else:
            for seg in speaker_segments:
                doc.add_paragraph(f"{seg['speaker']}: {seg['text']}")
        docx_path = filepath + '.docx'
        doc.save(docx_path)
        return jsonify({
            'success': True,
            'text': text_content,
            'download_url': f'/download?filename={filename}.txt',
            'download_docx': f'/download?filename={filename}.docx'
        })
    except Exception as e:
        return jsonify({'success': False, 'message': f'Error: {str(e)}'}), 500

@app.route('/download')
def download():
    filename = request.args.get('filename')
    if not filename:
        return 'Archivo no especificado', 400
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(filepath):
        return 'Archivo no encontrado', 404
    ext = os.path.splitext(filename)[1]
    download_name = 'transcripcion' + ext
    return send_file(filepath, as_attachment=True, download_name=download_name)

if __name__ == '__main__':
    app.run(debug=True)
